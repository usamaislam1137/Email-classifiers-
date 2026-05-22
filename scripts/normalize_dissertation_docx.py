#!/usr/bin/env python3
"""
Normalize Osama_Final_Combined_Dissertation (4).docx:
- Author line: Muhammad Usama ISLAM -> Muhammad Usama Islam (all occurrences; merge runs if needed)
- Black & white: all text/run colors black, no highlight; colored TABLE CELL fills (except cover
  table[0]) -> light gray. Cover page (first table) keeps green/brand colors. Images unchanged.
- Fonts: Times New Roman for body and headings (consistent sizes by heading level). Preserves
  Courier New blocks (code appendix).

Does not remove paragraphs, tables, or images. Creates a one-time backup before modifying.
"""
from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentType
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph

REPO = Path(__file__).resolve().parents[1]
DOC_PATH = REPO / "Osama_Final_Combined_Dissertation (4).docx"
COVER_TABLE_INDEX = 0

# Map saturated header / accent fills (outside cover) to print-safe gray
FILL_TO_GRAY = {
    "1F4E79": "E8E8E8",
    "FCEBD0": "F0F0F0",
}

CODE_BLOCK_FILL = "F4F6F9"  # keep appendix code shading as-is


def _backup_once(path: Path) -> Path:
    bak = path.with_name(path.stem + "_before_format_normalize.docx")
    if not bak.exists():
        shutil.copy2(path, bak)
        print("Backup:", bak)
    return bak


def _replace_text_in_paragraph(p: Paragraph, old: str, new: str) -> None:
    if old not in p.text:
        return
    for r in p.runs:
        if old in r.text:
            r.text = r.text.replace(old, new)
    if old not in p.text:
        return
    # Rare: string split across runs — merge into one run to complete replacement.
    merged = p.text.replace(old, new)
    for r in list(p.runs):
        p._p.remove(r._r)
    p.add_run(merged)


def _replace_in_every_paragraph(doc: DocumentType, old: str, new: str) -> None:
    for p in doc.paragraphs:
        _replace_text_in_paragraph(p, old, new)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_text_in_paragraph(p, old, new)
    for sec in doc.sections:
        for part in (sec.header, sec.footer, sec.first_page_header, sec.first_page_footer):
            try:
                if part is None:
                    continue
                for p in part.paragraphs:
                    _replace_text_in_paragraph(p, old, new)
                for table in part.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                _replace_text_in_paragraph(p, old, new)
            except Exception:
                continue


def _set_run_color_black(run) -> None:
    r_pr = run._r.get_or_add_rPr()
    for child in list(r_pr):
        if child.tag in (qn("w:color"), qn("w:highlight")):
            r_pr.remove(child)
    clr = OxmlElement("w:color")
    clr.set(qn("w:val"), "000000")
    r_pr.append(clr)


def _is_code_run(run) -> bool:
    n = run.font.name
    if not n:
        return False
    return "Courier" in n


def _heading_level(style_name: str) -> int:
    m = re.search(r"heading\s*(\d+)", (style_name or "Normal").lower())
    return int(m.group(1)) if m else 0


def _normalize_cell_fills_outside_cover(doc: DocumentType) -> None:
    for ti, table in enumerate(doc.tables):
        if ti == COVER_TABLE_INDEX:
            continue
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                tc_pr = tc.tcPr
                if tc_pr is None:
                    continue
                shd = tc_pr.find(qn("w:shd"))
                if shd is None:
                    continue
                fill = (shd.get(qn("w:fill")) or "").upper()
                if not fill or fill in ("FFFFFF", "AUTO", "NONE"):
                    continue
                if fill == CODE_BLOCK_FILL:
                    continue
                new_fill = FILL_TO_GRAY.get(fill, "EDEDED")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), new_fill)


def _paragraph_code_block_fill(p: Paragraph) -> str | None:
    p_pr = p._p.pPr
    if p_pr is None:
        return None
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        return None
    return (shd.get(qn("w:fill")) or "").upper()


def _normalize_runs_paragraph(p: Paragraph, *, allow_color: bool) -> None:
    if allow_color:
        return
    st = (p.style.name if p.style else "Normal") or "Normal"
    lvl = _heading_level(st)
    is_heading = lvl > 0 or st.lower().startswith("heading")

    for run in p.runs:
        if _is_code_run(run):
            _set_run_color_black(run)
            continue
        _set_run_color_black(run)
        run.font.name = "Times New Roman"
        if is_heading:
            sizes = {1: 16, 2: 14, 3: 13, 4: 12, 5: 11, 6: 11, 7: 11, 8: 11}
            run.font.size = Pt(sizes.get(lvl, 12) if lvl else 12)
            if lvl:
                run.bold = True
        elif st in ("List Paragraph", "TOC 1", "Table Paragraph", "Caption"):
            run.font.size = Pt(11)
        elif "toc" in st.lower():
            run.font.size = Pt(11)
        else:
            run.font.size = Pt(11)


def _walk_and_normalize_document_body(doc: DocumentType) -> None:
    for p in doc.paragraphs:
        fill = _paragraph_code_block_fill(p)
        if fill == CODE_BLOCK_FILL:
            for run in p.runs:
                _set_run_color_black(run)
            continue
        _normalize_runs_paragraph(p, allow_color=False)

    for ti, table in enumerate(doc.tables):
        allow = ti == COVER_TABLE_INDEX
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if allow:
                        continue
                    if _paragraph_code_block_fill(p) == CODE_BLOCK_FILL:
                        for run in p.runs:
                            _set_run_color_black(run)
                        continue
                    _normalize_runs_paragraph(p, allow_color=False)


def _normalize_headers_footers(doc: DocumentType) -> None:
    for sec in doc.sections:
        for part in (sec.header, sec.footer, sec.first_page_header, sec.first_page_footer):
            try:
                if part is None:
                    continue
                for p in part.paragraphs:
                    _normalize_runs_paragraph(p, allow_color=False)
                for table in part.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                _normalize_runs_paragraph(p, allow_color=False)
            except Exception:
                continue


def main() -> int:
    if not DOC_PATH.is_file():
        print("Missing", DOC_PATH)
        return 1
    _backup_once(DOC_PATH)
    doc = Document(str(DOC_PATH))

    _replace_in_every_paragraph(doc, "Muhammad Usama ISLAM", "Muhammad Usama Islam")
    _normalize_cell_fills_outside_cover(doc)
    _walk_and_normalize_document_body(doc)
    _normalize_headers_footers(doc)

    doc.save(str(DOC_PATH))
    print("Saved:", DOC_PATH)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
