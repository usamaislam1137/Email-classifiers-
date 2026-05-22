#!/usr/bin/env python3
"""
Remove runs of multiple empty body paragraphs that create huge blank gaps (common after
bad PDF/Word round-trips). Optionally set page_break_before on the following paragraph
for major headings only (Chapter, List of…), not before Figure/Table captions.

Backup once: Osama_Final_Combined_Dissertation (4)_before_gap_dedupe.docx
"""
from __future__ import annotations

import re
import shutil
import sys
from pathlib import Path

from docx import Document

REPO = Path(__file__).resolve().parents[1]
DOC_PATH = REPO / "Osama_Final_Combined_Dissertation (4).docx"


def _backup_once(path: Path) -> None:
    bak = path.with_name(path.stem + "_before_gap_dedupe.docx")
    if not bak.exists():
        shutil.copy2(path, bak)
        print("Backup:", bak)


def _delete_paragraph(paragraph) -> None:
    pel = paragraph._element
    parent = pel.getparent()
    parent.remove(pel)


def _find_empty_streaks(doc: Document) -> list[tuple[int, int, int]]:
    """Return (start_index, length, next_nonempty_index)."""
    paras = doc.paragraphs
    n = len(paras)
    out: list[tuple[int, int, int]] = []
    i = 0
    while i < n:
        if not (paras[i].text or "").strip():
            j = i
            while j < n and not (paras[j].text or "").strip():
                j += 1
            length = j - i
            if length >= 2:
                out.append((i, length, j if j < n else -1))
            i = j
        else:
            i += 1
    return out


def _needs_structural_page_break_after_streak(next_text: str) -> bool:
    t = (next_text or "").strip()
    if not t:
        return False
    if re.match(r"^Figure\s+\d", t, re.I):
        return False
    if re.match(r"^Table\s+\d", t, re.I):
        return False
    if t.upper().startswith("CHAPTER ") or t.startswith("Chapter "):
        return True
    if t.startswith("List of ") or t.upper().startswith("LIST OF "):
        return True
    if t.upper().startswith("REFERENCES") or t.upper().startswith("BIBLIOGRAPHY"):
        return True
    if t.upper().startswith("APPENDIX") or t.startswith("Appendix"):
        return True
    if t.upper().startswith("ABSTRACT"):
        return True
    return False


def main() -> int:
    if not DOC_PATH.is_file():
        print("Missing", DOC_PATH, file=sys.stderr)
        return 1
    _backup_once(DOC_PATH)
    doc = Document(str(DOC_PATH))

    streaks = _find_empty_streaks(doc)
    removed = 0
    page_breaks_set = 0

    for start, length, next_idx in sorted(streaks, key=lambda x: x[0], reverse=True):
        if length == 2:
            _delete_paragraph(doc.paragraphs[start + 1])
            removed += 1
            continue

        next_text = ""
        if 0 <= next_idx < len(doc.paragraphs):
            next_text = doc.paragraphs[next_idx].text
        for _ in range(length):
            _delete_paragraph(doc.paragraphs[start])
            removed += 1
        if start < len(doc.paragraphs) and _needs_structural_page_break_after_streak(next_text):
            doc.paragraphs[start].paragraph_format.page_break_before = True
            page_breaks_set += 1

    doc.save(str(DOC_PATH))
    print(
        f"Removed {removed} empty paragraph(s) from {len(streaks)} streaks; "
        f"set page_break_before on {page_breaks_set} structural heading(s). "
        f"Saved {DOC_PATH}"
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
