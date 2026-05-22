#!/usr/bin/env python3
"""
Append an code appendix to the dissertation Word file without editing earlier
paragraphs. Uses python-docx (append-only save).

Usage:
  python3 scripts/append_dissertation_appendix.py [--force]

--force  Append even if marker APPENDIX_EMAIL_PRIORITY_CODE_V2 already exists.
"""
from __future__ import annotations

import argparse
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

MARKER = "APPENDIX_EMAIL_PRIORITY_CODE_V2"
REPO = Path(__file__).resolve().parents[1]
DOC_PATH = REPO / "Osama_Final_Combined_Dissertation (4).docx"
EPS = REPO / "email_priority_system"

SECTIONS: list[tuple[str, Path | None, str]] = [
    (
        "Note on layout and CodeSnap",
        None,
        f"This appendix ({MARKER}) was generated from the project repository; the "
        "main chapters above were not modified. Each listing is the actual source text. "
        "For submission, you may add figures taken with the CodeSnap extension (VS Code / "
        "Cursor): open the file, select the code, run CodeSnap, export PNG, and paste the "
        "image in Word above the corresponding listing for readability. "
        "Each listing below is formatted as a shaded code block (background + "
        "accent border) for readability.",
    ),
    (
        "ML — configuration",
        EPS / "ml/config.py",
        "Paths, labels, class thresholds, training limits, Flask settings, and keyword corpora.",
    ),
    (
        "ML — dataset download",
        EPS / "ml/download_dataset.py",
        "Downloads and extracts Enron and SpamAssassin corpora used for training.",
    ),
    (
        "ML — preprocessing",
        EPS / "ml/preprocess.py",
        "Parsing, cleaning, and semi-supervised priority labelling of raw mail.",
    ),
    (
        "ML — feature engineering",
        EPS / "ml/feature_engineering.py",
        "TF-IDF, metadata features, and optional BERT embeddings.",
    ),
    (
        "ML — model training",
        EPS / "ml/train_models.py",
        "Trains logistic regression, random forest, XGBoost, and optional DistilBERT.",
    ),
    (
        "ML — evaluation and model selection",
        EPS / "ml/evaluate_models.py",
        "Metrics, best-model selection, SHAP, and evaluation report output.",
    ),
    (
        "ML — rule-based fallback",
        EPS / "ml/fallback_model.py",
        "Deterministic fallback when learned models fail quality thresholds.",
    ),
    (
        "ML — single-email prediction",
        EPS / "ml/predict.py",
        "loads the chosen model and returns prediction + confidence for the API.",
    ),
    (
        "ML — REST API (Flask)",
        EPS / "ml/flask_api.py",
        "HTTP endpoints: /predict, /health, /model_info, batch predict, etc.",
    ),
    (
        "ML — pipeline shell driver",
        EPS / "ml/run_pipeline.sh",
        "Single-command orchestration of download → train → evaluate for reproducibility.",
    ),
    (
        "Rails — ML API client",
        EPS / "rails_app/app/services/ml_api_client.rb",
        "HTTP client used by the web app to call the Python classifier service.",
    ),
    (
        "Rails — routes",
        EPS / "rails_app/config/routes.rb",
        "URL mapping for dashboard, classifications, and history.",
    ),
    (
        "Rails — dashboard controller",
        EPS / "rails_app/app/controllers/dashboard_controller.rb",
        "Overview page and aggregate stats.",
    ),
    (
        "Rails — classifications controller",
        EPS / "rails_app/app/controllers/classifications_controller.rb",
        "Form submit, ML API call, persistence, and result display.",
    ),
]


def _apply_code_block_paragraph_style(paragraph, *, fill_hex: str = "F4F6F9") -> None:
    """Shaded background, subtle box border, IDE-style left accent — reads as a code snippet in Word."""
    p_pr = paragraph._p.get_or_add_pPr()

    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "80")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "240")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)

    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "200")
    ind.set(qn("w:right"), "200")
    p_pr.append(ind)

    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    p_pr.append(shd)

    p_bdr = OxmlElement("w:pBdr")
    thin = ("4", "D0D7DE")  # sz (eighths of pt), color
    accent = ("28", "0969DA")
    for side, (sz, color) in (
        ("top", thin),
        ("right", thin),
        ("bottom", thin),
        ("left", accent),
    ):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "6")
        el.set(qn("w:color"), color)
        p_bdr.append(el)
    p_pr.append(p_bdr)


def _mono_paragraph(doc: Document, text: str, size_pt: int = 9) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(size_pt)
    _apply_code_block_paragraph_style(p)


def _body_msg(doc: Document, text: str, *, italic: bool = False, size: int = 11) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"


def _section_title(doc: Document, text: str, *, level: int = 1) -> None:
    """Word files from some templates lack built-in Heading 1/2 styles; use explicit formatting."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(16 if level == 1 else 13)


def _add_code(doc: Document, content: str, chunk_lines: int = 55) -> None:
    lines = content.splitlines()
    for i in range(0, len(lines), chunk_lines):
        chunk = "\n".join(lines[i : i + chunk_lines])
        _mono_paragraph(doc, chunk)


def _doc_plain_text(doc: Document) -> str:
    return "\n".join(p.text for p in doc.paragraphs)


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--force", action="store_true")
    args = ap.parse_args()

    if not DOC_PATH.is_file():
        print(f"Missing document: {DOC_PATH}", file=sys.stderr)
        return 1

    backup = DOC_PATH.with_name(DOC_PATH.stem + "_backup_before_code_appendix.docx")
    if not backup.exists():
        shutil.copy2(DOC_PATH, backup)
        print(f"Backup created: {backup}")

    doc = Document(str(DOC_PATH))
    body = _doc_plain_text(doc)
    if MARKER in body and not args.force:
        print(
            f"Appendix marker already present ({MARKER}). "
            "Open Word, remove the appendix if you want to regenerate, or pass --force.",
            file=sys.stderr,
        )
        return 2

    doc.add_page_break()
    _section_title(doc, "Appendix: Source code listings (Email priority system)", level=1)

    for heading, path, blurb in SECTIONS:
        if path is None:
            _body_msg(doc, blurb, italic=False, size=11)
            continue
        if not path.is_file():
            print(f"Skip missing file: {path}", file=sys.stderr)
            continue
        _section_title(doc, heading, level=2)
        _body_msg(doc, blurb, italic=True, size=11)
        _body_msg(
            doc,
            "Optional — CodeSnap: capture this file in the IDE and paste the image here.",
            italic=True,
            size=9,
        )
        raw = path.read_text(encoding="utf-8", errors="replace")
        _body_msg(doc, f"File: {path.relative_to(REPO)}", size=9)
        _add_code(doc, raw)
        doc.add_paragraph()

    doc.save(str(DOC_PATH))
    print(f"Updated (appendix only): {DOC_PATH}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
